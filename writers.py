# Module containing functions that write analysed data to various file types

# Import modules

from csv import writer as csv_writer, QUOTE_ALL
from re import search, findall, sub
from docx import Document  # Version 1.2.0
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def nmr_csv_writer(analysed_data):
    """Creates a CSV file with Boltzmann-averaged NMR shielding tensors (and if calculated,
    scaled chemical shifts too)."""

    # Define input data

    c_nmr = analysed_data[11]
    h_nmr = analysed_data[12]
    results_name_and_directory = analysed_data[27]
    # Write CSV file

    csv_name = results_name_and_directory + " Boltzmann-Averaged NMR Data.csv"
    title = ["Boltzmann-Averaged NMR Data"]
    column_headings = ["Element & Atom Number", "Element & Number in Group", "Shielding tensor"]
    if len(c_nmr[0]) == 4 or len(h_nmr[0]) == 4:  # Detect if have chemical shifts too
        column_headings.append("Scaled chemical shift (ppm)")
    try:
        with open(csv_name, "w", newline="") as file:
            writer = csv_writer(file, quoting=QUOTE_ALL)
            writer.writerow(title)
            writer.writerow(column_headings)
            for carbon in range(len(c_nmr)):
                writer.writerow(c_nmr[carbon])
            writer.writerow("")
            for hydrogen in range(len(h_nmr)):
                writer.writerow(h_nmr[hydrogen])
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save NMR data .CSV file.\n Try closing this .CSV file, then drag and drop" " your files again.",
        )


def ir_csv_writer(analysed_data, settings):
    """Creates a CSV file containing IR data - raw and Boltzmann-averaged (scaled) vibrational frequencies
    and IR intensities."""

    # Define input data

    boltz_frequencies = analysed_data[29]
    ordered_frequencies = analysed_data[30]
    boltz_ir_intensities = analysed_data[51]
    
    ordered_ir_intensities = analysed_data[52]
    ordered_boltz_weights = analysed_data[2]
    results_name_and_directory = analysed_data[27]

    # Compile data into rows for CSV file

    boltz_percent = []
    for i in ordered_boltz_weights:
        boltz_percent.append(str(i * 100) + "%")
        boltz_percent.append("")
    ordered_ir = []
    for freq in range(len(boltz_frequencies)):
        line = []
        for conf in range(len(ordered_frequencies)):
            line.append(ordered_frequencies[conf][freq])
            line.append(ordered_ir_intensities[conf][freq])
        ordered_ir.append(line)
    data = []
    if settings["IR freq scaling factor"] == "":
        for i in range(-1, len(boltz_frequencies)):
            if i == -1:
                data.append(["Unscaled", "", "", "", "", "", "", "", "", "", ""] + boltz_percent)
            else:
                data.append(
                    [boltz_frequencies[i], "", "", "", "", "", boltz_ir_intensities[i], "", "", "", ""] + ordered_ir[i]
                )
    elif settings["IR freq scaling factor"] != "":
        for i in range(-1, len(boltz_frequencies)):
            if i == -1:
                data.append(
                    [
                        "Unscaled",
                        "Scaled (scaling factor = " + settings["IR freq scaling factor"] + ")",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                        "",
                    ]
                    + boltz_percent
                )
            else:
                data.append(
                    [
                        boltz_frequencies[i],
                        str(boltz_frequencies[i] * float(settings["IR freq scaling factor"])),
                        "",
                        "",
                        "",
                        "",
                        boltz_ir_intensities[i],
                        "",
                        "",
                        "",
                    ]
                    + ordered_ir[i]
                )
    # Write CSV file

    csv_name = results_name_and_directory + " IR Data.csv"
    title = [
        "Boltzmann-Averaged Vibrational Frequencies (cm^-1)",
        "",
        "",
        "",
        "",
        "",
        "Boltzmann-Averaged IR Intensities (km/mol)",
        "",
        "",
        "",
        "",
        "Conformer Boltzmann populations, unscaled vibrational frequencies (cm^-1) and IR intensities (km/mol), arranged in order of decreasing Boltzmann population",
    ]
    try:
        with open(csv_name, "w", newline="") as file:
            writer = csv_writer(file, quoting=QUOTE_ALL)
            writer.writerow(title)
            for line in data:
                writer.writerow(line)
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save frequency data .CSV file.\n Try closing this .CSV file, then drag and drop your files again.",
        )


def cd_bil_writer(analysed_data):
    """Creates a .cd.bil file containing the Boltzmann-averaged excited state CD data, for SpecDis"""

    # Define input data

    boltz_wavelengths = analysed_data[13]
    boltz_rotatory_strengths = analysed_data[14]
    results_name_and_directory = analysed_data[27]
    ecd_file_name = results_name_and_directory + " Boltzmann-Averaged ECD Data.cd.bil"
    # Write .cd.bil file

    try:
        with open(ecd_file_name, "w") as file:
            file.write(
                "SpecDis bil-file (length formalism) wavelength (nm)       rotational strength (10**-40 esu**2-cm**2)"
            )
            for excited_state in range((len(boltz_wavelengths) - 1), -1, -1):
                file.write(
                    "\n" + str(boltz_wavelengths[excited_state]) + "\t" + str(boltz_rotatory_strengths[excited_state])
                )
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save ECD data .cd.bil file.\n Try closing this .cd.bil file, then drag and drop your files again.",
        )


def uv_bil_writer(analysed_data):
    """Creates a .uv.bil file containing the Boltzmann-averaged excited state UV data, for SpecDis"""

    # Define input data

    boltz_wavelengths = analysed_data[13]
    boltz_oscillator_strengths = analysed_data[28]
    results_name_and_directory = analysed_data[27]
    uv_file_name = results_name_and_directory + " Boltzmann-Averaged UV Data.uv.bil"
    # Write .uv.bil file

    try:
        with open(uv_file_name, "w") as file:
            file.write("SpecDis bil-file (length formalism) wavelength (nm)       oscillator strength")
            for excited_state in range((len(boltz_wavelengths) - 1), -1, -1):
                file.write(
                    "\n" + str(boltz_wavelengths[excited_state]) + "\t" + str(boltz_oscillator_strengths[excited_state])
                )
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save UV data .uv.bil file.\n Try closing this .uv.bil file, then drag and drop your files again.",
        )


def ir_bil_writer(analysed_data):
    """Creates a .ir.bil file containing Boltzmann-averaged IR data, for SpecDis"""

    # Define input data

    boltz_frequencies = analysed_data[29]
    boltz_frequency_dipole_strengths = analysed_data[40]
    results_name_and_directory = analysed_data[27]
    calc_software = analysed_data[55]
    ir_file_name = results_name_and_directory + " Boltzmann-Averaged IR Data.ir.bil"
    # Write .ir.bil file

    try:
        with open(ir_file_name, "w") as file:
            if calc_software == "gaussian":
                file.write("SpecDis bil-file (wavenumber (cm**-1)           dipole strength (10**-40 esu**2-cm**2))")
            elif calc_software == "orca":
                file.write("SpecDis bil-file (wavenumber (cm**-1)           molar extinction coefficient (L mol**-1 cm**-1)")
            for frequency in range((len(boltz_frequencies) - 1), -1, -1):
                file.write(
                    "\n" + str(boltz_frequencies[frequency]) + "\t" + str(boltz_frequency_dipole_strengths[frequency])
                )
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save IR data .ir.bil file.\n Try closing this .ir.bil file, then drag and drop your files again.",
        )


def vc_bil_writer(analysed_data):
    """Creates a .vc.bil file containing the Boltzmann-averaged VCD data, for SpecDis"""

    # Define input data

    boltz_frequencies = analysed_data[29]
    boltz_frequency_rotatory_strengths = analysed_data[35]
    results_name_and_directory = analysed_data[27]
    vc_file_name = results_name_and_directory + " Boltzmann-Averaged VCD Data.vc.bil"
    # Write .vc.bil file

    try:
        with open(vc_file_name, "w") as file:
            file.write("SpecDis bil-file VCD (wavenumber (cm**-1)   rotational strength (10**-44 esu**2-cm**2))")
            for frequency in range((len(boltz_frequencies) - 1), -1, -1):
                file.write(
                    "\n" + str(boltz_frequencies[frequency]) + "\t" + str(boltz_frequency_rotatory_strengths[frequency])
                )
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save vcd data .vc.bil file.\n Try closing this .vc.bil file, then drag and drop your files again.",
        )


def or_bil_writer(analysed_data):
    """Creates a .or.bil file containing the Boltzmann-averaged optical rotation data, for SpecDis"""

    # Define input data

    ordered_optrot_wavelengths_list = analysed_data[36]
    boltz_optrot_strengths = analysed_data[38]
    results_name_and_directory = analysed_data[27]
    or_file_name = results_name_and_directory + " Boltzmann-Averaged OR Data.or.bil"
    # Write .or.bil file

    try:
        with open(or_file_name, "w") as file:
            file.write("SpecDis bil-file (length formalism) wavelength (nm)       oscillator strength")
            for wavelength in range((len(ordered_optrot_wavelengths_list[0]) - 1), -1, -1):
                file.write(
                    "\n"
                    + str(ordered_optrot_wavelengths_list[0][wavelength])
                    + "\t"
                    + str(boltz_optrot_strengths[wavelength])
                )
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save optical rotation data .or.bil file.\n Try closing this .or.bil file, then drag and drop your files again.",
        )


def xyz_writer(analysed_data):
    """Writes an XYZ file containing conformer geometries and energies, arranged in order of increasing energy."""

    # Define input data

    energies = analysed_data[0]
    element_list = analysed_data[3]
    x_cartesian_coords_list = analysed_data[4]
    y_cartesian_coords_list = analysed_data[5]
    z_cartesian_coords_list = analysed_data[6]
    results_name_and_directory = analysed_data[27]
    xyz_file_name = results_name_and_directory + " Geometries and Energies.xyz"
    # Write .xyz file

    try:
        with open(xyz_file_name, "w") as file:
            for conformer in range(len(energies)):
                file.write(str(len(x_cartesian_coords_list[0])) + "\n" + str(energies[conformer]) + "\n")
                for atom in range(len(element_list[conformer])):
                    x = f"{float(x_cartesian_coords_list[conformer][atom]):.10f}"
                    space1 = (
                        " " * (2 - len(element_list[conformer][atom]))
                        + " " * (7 - len(str(int(float(x)))))
                        + " " * (1 - x.count("-0."))
                    )
                    y = f"{float(y_cartesian_coords_list[conformer][atom]):.10f}"
                    space2 = " " * (9 - len(str(int(float(y))))) + " " * (1 - y.count("-0."))
                    z = f"{float(z_cartesian_coords_list[conformer][atom]):.10f}"
                    space3 = " " * (9 - len(str(int(float(z))))) + " " * (1 - z.count("-0."))
                    file.write(element_list[conformer][atom] + space1 + x + space2 + y + space3 + z + "\n")
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save .xyz file.\n Try closing this .xyz file, then drag and drop your files again. ",
        )


def dup_conf_txt_writer(analysed_data, settings):
    """Creates a .txt file containing the energy differences and MADs of redundant conformers."""

    # Define input data

    if len(analysed_data) == 6:  # This means data is from xyz/sdf files
        duplicate_conf_details = analysed_data[4]
        results_name_and_directory = analysed_data[3]
        numbering_message = ""
        thresholds_message = (
            "Redundant conformer detection threshold:\nMAD = " + str(settings["MAD cutoff (A)"]) + " angstroms\n\n "
        )
    else:  # This means data is from .out/.log files
        if not analysed_data[50]:  # This means conformer suffixes are not present
            numbering_message = "Conformers are numbered by their order of appearance in the selected files"
            if analysed_data[53] > 0:  # This means conformers with imaginary frequencies were removed
                numbering_message += ". Conformers with imaginary frequencies are excluded.\n"
            else:
                numbering_message += ".\n"
        else:
            numbering_message = ""
        duplicate_conf_details = analysed_data[44]
        results_name_and_directory = analysed_data[27]
        thresholds_message = (
            "Redundant conformer detection thresholds:\nEnergy difference = "
            + str(settings["Energy cutoff (kcal/mol)"])
            + " kcal/mol\nMAD = "
            + str(settings["MAD cutoff (A)"])
            + " angstroms.\n\n"
        )
    dup_conf_text = ""
    for text in duplicate_conf_details:
        dup_conf_text += text + "\n"
    if len(duplicate_conf_details) > 1:
        dup_conf_txt_file_name = results_name_and_directory + " Redundant Conformers.txt"
    else:
        dup_conf_txt_file_name = results_name_and_directory + " Redundant Conformer.txt"
    # Write .txt file

    try:
        with open(dup_conf_txt_file_name, "w", encoding="utf-8") as file:
            file.write(numbering_message)
            file.write(thresholds_message)
            file.write(dup_conf_text)
        file.close()
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save redundant conformer .txt file.\n Try closing this .txt file, "
            "then drag and drop your files again. ",
        )


def input_file_writer(data, settings, filename):
    """Writes quantum chemistry calculation input files (e.g. for Gaussian or ORCA)."""

    # Define input data

    removals = 0
    if len(data) == 6:  # This means data is from xyz/sdf files
        element_list = data[0][0]
        x_cartesian_coords_list = data[0][1]
        y_cartesian_coords_list = data[0][2]
        z_cartesian_coords_list = data[0][3]
        results_name_and_directory = data[3]
        directory = sub(r"/.[^/]*$", r"/", results_name_and_directory)
        results_name_and_directory = directory + filename
        conformer_numbers_list = []
        for i in range(1, (len(x_cartesian_coords_list) + 1)):
            conformer_numbers_list.append(str(i))
    else:  # This means data is from .out/.log files
        if not settings["Geometry reopt relative energy threshold"].replace(".", "", 1).isdigit():  # Check for error
            parser_error_check = "Error detected"
            error_message = "Conformer energy window setting contains a non-number value."
            return parser_error_check, error_message
        rel_energy_threshold = float(settings["Geometry reopt relative energy threshold"])
        rel_energy_threshold_unit = settings["Geometry reopt relative energy unit"]
        relative_energies = data[45]
        element_list = data[46]
        x_cartesian_coords_list = data[47]
        y_cartesian_coords_list = data[48]
        z_cartesian_coords_list = data[49]
        if data[54]:
            pre_conformer_suffixes_list = data[54]
        else:
            pre_conformer_suffixes_list = data[50]
        conformer_numbers_list = []
        for i in range(0, len(pre_conformer_suffixes_list)):
            conformer_suffix = pre_conformer_suffixes_list[i].removesuffix(".log").removesuffix(".out")
            conformer_numbers_list.append(findall(r"\d+", conformer_suffix)[0])
        if not conformer_numbers_list:
            for i in range(1, (len(element_list) + 1)):
                conformer_numbers_list.append(i)
        results_name_and_directory = data[27]
        directory = sub(r"/.[^/]*$", r"/", results_name_and_directory)
        results_name_and_directory = directory + filename
        # Remove conformers above relative energy threshold

        if (
            settings["Geometry reopt"] is True
            and settings["Skip excluding duplicate conformers from input files made from output files"] is False
        ):
            if str(rel_energy_threshold).endswith(".0"):
                rel_energy_threshold = int(rel_energy_threshold)
            results_name_and_directory += "_" + str(rel_energy_threshold) + rel_energy_threshold_unit
            # If needed, convert the relative energy threshold units, to the units already used to calculate relative
            # energies

            if rel_energy_threshold_unit == "kJmol-1" and settings["Relative energy unit"] == "kcal/mol":
                rel_energy_threshold = rel_energy_threshold / 4.184
            elif rel_energy_threshold_unit == "kcalmol-1" and settings["Relative energy unit"] == "kJ/mol":
                rel_energy_threshold = rel_energy_threshold * 4.184
            for conf_number, energy in enumerate(relative_energies):
                conf_number = conf_number - removals
                if energy > rel_energy_threshold:
                    # Delete conformer data

                    del element_list[conf_number]
                    del x_cartesian_coords_list[conf_number]
                    del y_cartesian_coords_list[conf_number]
                    del z_cartesian_coords_list[conf_number]
                    del conformer_numbers_list[conf_number]
                    removals += 1  # This is because indices of remaining conformers are now reduced by 1,
                    # by removing the last conformer
    input_files = []
    if settings["Input File Conformers Together"] is True:
        input_files.append(results_name_and_directory + ".inp")
    if settings["Input File Conformers Together"] is False:
        for i in range(0, len(conformer_numbers_list)):  # iterate through all (remaining) conformers
            input_files.append(results_name_and_directory + "_conf-" + str(conformer_numbers_list[i]) + ".inp")
    # Write input file(s)

    conf_tracker = 0
    try:
        for inp_file in input_files:
            with open(inp_file, "w", encoding="utf-8") as file:
                for index, conformer_number in enumerate(
                    conformer_numbers_list
                ):  # iterate through all (remaining) conformers
                    # Create atom coords block

                    atom_coords_block = ""
                    for atom in range(len(element_list[index])):
                        element = element_list[index][atom]
                        x = f"{float(x_cartesian_coords_list[index][atom]):.10f}"
                        space1 = "  " + " " * (1 - x.count("-"))
                        if len(element) > 1:
                            space1 = " " + " " * (1 - x.count("-"))
                        y = f"{float(y_cartesian_coords_list[index][atom]):.10f}"
                        space2 = "  " + " " * (1 - y.count("-"))
                        z = f"{float(z_cartesian_coords_list[index][atom]):.10f}"
                        space3 = "  " + " " * (1 - z.count("-"))
                        atom_coords_block += element + space1 + x + space2 + y + space3 + z + "\n"
                    if settings["Input File Conformers Together"] is False and index < conf_tracker:
                        continue
                    if settings["Input File Conformers Together"] is False and index > conf_tracker:
                        break
                    for i, text in enumerate(
                        settings["Input File Texts"]
                    ):  # iterate through calculations for each conformer
                        text = text.replace("⫷⫷⫷CONFORMER NUMBER⫸⫸⫸", str(conformer_number))
                        text = text.replace("⫷⫷⫷COMPOUND NAME⫸⫸⫸", filename)
                        text = text.replace("⫷⫷⫷ATOM COORDINATES⫸⫸⫸", atom_coords_block)
                        if (
                            i == 0 and index != 0 and settings["Input File Conformers Together"] is True
                        ):  # multi job text
                            file.write(text + "\n")
                        if i > 1:
                            file.write(settings["Input File Texts"][0] + "\n")  # multi job text
                        if i > 0:
                            file.write(text + "\n")
                file.write("\n\n")  # Some comp chem software (Gaussian) require some blank lines at end of input file
            file.close()
            conf_tracker += 1
        return "", "", removals, len(input_files)
    except:
        return (
            "Error detected",
            "Unable to save .inp file.\n Try closing this .inp file, then drag and drop your files again.",
        )


def docx_writer(analysed_data, settings):
    """Creates a nicely formatted Microsoft Word document (.docx) containing supplementary data tables,
    suitable for publication."""

    # Return no docx file if neither an energies/coords table or NMR/ECD/VCD/OR table are requested

    if settings["Energies and coordinates table"] is False and settings["NMR/ECD/VCD/OR table"] is False:
        return "", ""
    # Define input data

    ordered_energies = analysed_data[0]
    ordered_relative_energies = analysed_data[1]
    ordered_boltz_weights = analysed_data[2]
    ordered_element_list = analysed_data[3]
    ordered_x_cartesian_coords_list = analysed_data[4]
    ordered_y_cartesian_coords_list = analysed_data[5]
    ordered_z_cartesian_coords_list = analysed_data[6]
    ordered_wavelength_list = analysed_data[7]
    ordered_rotatory_strength_list = analysed_data[8]
    ordered_shielding_tensors = analysed_data[9]
    sp_functional_and_basis_set = analysed_data[15]
    sp_solvent = analysed_data[16]
    opt_freq_functional_and_basis_set = analysed_data[18]
    opt_freq_solvent = analysed_data[19]
    nmr_functional_and_basis_set = analysed_data[21]
    nmr_solvent = analysed_data[22]
    tddft_functional_and_basis_set = analysed_data[24]
    tddft_solvent = analysed_data[25]
    results_name_and_directory = analysed_data[27]
    ordered_frequencies = analysed_data[30]
    ordered_oscillator_strengths = analysed_data[31]
    ordered_frequency_rotatory_strengths_list = analysed_data[34]
    ordered_optrot_wavelengths_list = analysed_data[36]
    ordered_optrot_list = analysed_data[37]
    ordered_frequency_dipole_strengths_list = analysed_data[39]
    or_functional_and_basis_set = analysed_data[41]
    or_solvent = analysed_data[42]
    doc_text = ""

    # Define new lists

    four_conformers_energies = []
    four_conformers_relative_energies = []
    four_conformers_boltz_weights = []
    four_conformers_element_list = []
    four_conformers_x_cartesian_coords_list = []
    four_conformers_y_cartesian_coords_list = []
    four_conformers_z_cartesian_coords_list = []
    four_conformers_wavelength_list = []
    four_conformers_rotatory_strength_list = []
    four_conformers_frequencies = []
    four_conformers_frequency_rotatory_strengths = []
    four_conformers_oscillator_strength_list = []
    four_conformers_frequency_dipole_strengths = []
    six_conformers_shielding_tensors = []
    six_conformers_optrots = []
    table_energies = []
    table_relative_energies = []
    table_boltz_weights = []
    table_element_list = []
    table_x_cartesian_coords_list = []
    table_y_cartesian_coords_list = []
    table_z_cartesian_coords_list = []
    table_wavelength_list = []
    table_rotatory_strength_list = []
    table_oscillator_strength_list = []
    table_shielding_tensors = []
    table_frequencies = []
    table_frequency_rotatory_strengths = []
    table_frequency_dipole_strengths = []
    table_optrots = []

    # Create name for compound in Word document

    results_name = search("/([^/]+)$", results_name_and_directory).group(1)

    # Organise non-NMR/ECD/VCD data into groups containing up to four conformers each
    # (to fit nicely into a Word table, up to four conformers wide)

    j = 0
    for conformer_number, conformer in enumerate(ordered_energies):
        four_conformers_energies.append(ordered_energies[conformer_number])
        four_conformers_relative_energies.append(ordered_relative_energies[conformer_number])
        four_conformers_boltz_weights.append(ordered_boltz_weights[conformer_number])
        four_conformers_element_list.append(ordered_element_list[conformer_number])
        four_conformers_x_cartesian_coords_list.append(ordered_x_cartesian_coords_list[conformer_number])
        four_conformers_y_cartesian_coords_list.append(ordered_y_cartesian_coords_list[conformer_number])
        four_conformers_z_cartesian_coords_list.append(ordered_z_cartesian_coords_list[conformer_number])
        if j < 3:
            j += 1
        elif j == 3:
            table_energies.append(four_conformers_energies)
            table_relative_energies.append(four_conformers_relative_energies)
            table_boltz_weights.append(four_conformers_boltz_weights)
            table_element_list.append(four_conformers_element_list)
            table_x_cartesian_coords_list.append(four_conformers_x_cartesian_coords_list)
            table_y_cartesian_coords_list.append(four_conformers_y_cartesian_coords_list)
            table_z_cartesian_coords_list.append(four_conformers_z_cartesian_coords_list)
            four_conformers_energies = []
            four_conformers_relative_energies = []
            four_conformers_boltz_weights = []
            four_conformers_element_list = []
            four_conformers_x_cartesian_coords_list = []
            four_conformers_y_cartesian_coords_list = []
            four_conformers_z_cartesian_coords_list = []
            j = 0
    if j > 0:
        table_energies.append(four_conformers_energies)
        table_relative_energies.append(four_conformers_relative_energies)
        table_boltz_weights.append(four_conformers_boltz_weights)
        table_element_list.append(four_conformers_element_list)
        table_x_cartesian_coords_list.append(four_conformers_x_cartesian_coords_list)
        table_y_cartesian_coords_list.append(four_conformers_y_cartesian_coords_list)
        table_z_cartesian_coords_list.append(four_conformers_z_cartesian_coords_list)
    # Organise ECD data into groups containing up to four conformers each
    # (to fit nicely into a Word table, up to four conformers wide)

    if ordered_wavelength_list:
        k = 0
        for conformer_number, conformer in enumerate(ordered_energies):
            four_conformers_wavelength_list.append(ordered_wavelength_list[conformer_number])
            four_conformers_rotatory_strength_list.append(ordered_rotatory_strength_list[conformer_number])
            four_conformers_oscillator_strength_list.append(ordered_oscillator_strengths[conformer_number])
            if k < 3:
                k += 1
            elif k == 3:
                table_wavelength_list.append(four_conformers_wavelength_list)
                table_rotatory_strength_list.append(four_conformers_rotatory_strength_list)
                table_oscillator_strength_list.append(four_conformers_oscillator_strength_list)
                four_conformers_wavelength_list = []
                four_conformers_rotatory_strength_list = []
                four_conformers_oscillator_strength_list = []
                k = 0
        if k > 0:
            table_wavelength_list.append(four_conformers_wavelength_list)
            table_rotatory_strength_list.append(four_conformers_rotatory_strength_list)
            table_oscillator_strength_list.append(four_conformers_oscillator_strength_list)
    # Organise VCD data into groups containing up to four conformers each
    # (to fit nicely into a Word table, up to four conformers wide)

    if ordered_frequency_rotatory_strengths_list:
        k = 0
        for conformer_number, conformer in enumerate(ordered_energies):
            four_conformers_frequencies.append(ordered_frequencies[conformer_number])
            four_conformers_frequency_rotatory_strengths.append(
                ordered_frequency_rotatory_strengths_list[conformer_number]
            )
            four_conformers_frequency_dipole_strengths.append(ordered_frequency_dipole_strengths_list[conformer_number])
            if k < 3:
                k += 1
            elif k == 3:
                table_frequencies.append(four_conformers_frequencies)
                table_frequency_rotatory_strengths.append(four_conformers_frequency_rotatory_strengths)
                table_frequency_dipole_strengths.append(four_conformers_frequency_dipole_strengths)
                four_conformers_frequencies = []
                four_conformers_frequency_rotatory_strengths = []
                four_conformers_frequency_dipole_strengths = []
                k = 0
        if k > 0:
            table_frequencies.append(four_conformers_frequencies)
            table_frequency_rotatory_strengths.append(four_conformers_frequency_rotatory_strengths)
            table_frequency_dipole_strengths.append(four_conformers_frequency_dipole_strengths)
    # Organise NMR data into groups containing up to six conformers each
    # (to fit nicely into a Word table, up to six conformers wide)

    if ordered_shielding_tensors:
        l = 0
        for conformer_number, conformer in enumerate(ordered_energies):
            six_conformers_shielding_tensors.append(ordered_shielding_tensors[conformer_number])
            if l < 5:
                l += 1
            elif l == 5:
                table_shielding_tensors.append(six_conformers_shielding_tensors)
                six_conformers_shielding_tensors = []
                l = 0
        if l > 0:
            table_shielding_tensors.append(six_conformers_shielding_tensors)
    if ordered_optrot_wavelengths_list:
        # Add positive signs and degrees symbol to optrot strengths

        new_list = []
        for wavelengths in ordered_optrot_list:
            new_list2 = []
            for optrot_strength in wavelengths:
                if float(optrot_strength) > 0:
                    optrot_strength = "+" + optrot_strength
                optrot_strength = optrot_strength + "º"
                new_list2.append(optrot_strength)
            new_list.append(new_list2)
        ordered_optrot_list = new_list
        # Organise optrot data into groups containing up to six conformers each
        # (to fit nicely into a Word table, up to six conformers wide)

        l = 0
        for conformer_number, conformer in enumerate(ordered_energies):
            six_conformers_optrots.append(ordered_optrot_list[conformer_number])
            if l < 5:
                l += 1
            elif l == 5:
                table_optrots.append(six_conformers_optrots)
                six_conformers_optrots = []
                l = 0
        if l > 0:
            table_optrots.append(six_conformers_optrots)

    def set_cell_border(cell, **kwargs):
        """
        Set a specified cell`s border in Microsoft Word tables

        This function is from https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx by MadisonTrash
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Check for tag existence, if none found, then create one

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        # List over all available tags

        for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = "w:{}".format(edge)

                # Check for tag existence, if none found, then create one

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                # Looks like order of attributes is important

                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn("w:{}".format(key)), str(edge_data[key]))

    def set_cell_margins(table, left=0.1, right=0.1):
        """Function to set cell margins in Microsoft Word tables.

        This function is from https://stackoverflow.com/questions/51060431/how-to-set-cell-margins-of-tables-in-ms-word-using-python-docx by mtytgat
        """
        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement("w:tblCellMar")
        kwargs = {"left": left, "right": right}
        for m in ["left", "right"]:
            node = OxmlElement("w:{}".format(m))
            node.set(qn("w:w"), str(kwargs.get(m)))
            node.set(qn("w:type"), "dxa")
            tblCellMar.append(node)
        tblPr.append(tblCellMar)

    if settings["Energies and coordinates table"] is True:
        # Create table title in Word document

        if sp_functional_and_basis_set:
            if not opt_freq_solvent:
                opt_freq_solvation_text = " in the gas phase"
            else:
                opt_freq_solvation_text = " in " + opt_freq_solvent
            if not sp_solvent:
                sp_solvation_text = " in the gas phase"
            else:
                sp_solvation_text = " in " + sp_solvent
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            table_title = doc.add_heading("", level=2)
            table_title.add_run(
                "Energies and atomic Cartesian coordinates of geometry-optimized conformers of "
                + results_name
                + " at "
                + opt_freq_functional_and_basis_set
                + opt_freq_solvation_text
                + ", with electronic energies at "
                + sp_functional_and_basis_set
                + sp_solvation_text
                + "\n"
            )
        else:
            if not opt_freq_solvent:
                opt_freq_solvation_text = " in the gas phase"
            else:
                opt_freq_solvation_text = " in " + opt_freq_solvent
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            table_title = doc.add_heading("", level=2)
            table_title.add_run(
                "Energies and atomic Cartesian coordinates of geometry-optimized conformers of "
                + results_name
                + " at "
                + opt_freq_functional_and_basis_set
                + opt_freq_solvation_text
                + "\n"
            )
        # Create style

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1

        number_of_columns = 1 + 3 * len(table_energies[0])
        # Create table

        for four_conformers in range(0, len(table_energies)):
            table = doc.add_table(1, number_of_columns)
            set_cell_margins(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Remove extra first row if is concatenated table

            if four_conformers > 0:
                Table = doc.tables[four_conformers]
                RowA = Table.rows[0]
                table_element = Table._tbl
                table_element.remove(RowA._tr)
            if len(table_energies) > 1:
                table_loop = range(4)
            else:
                table_loop = range(len(table_energies[four_conformers]))
            # Add conformer numbers

            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                if four_conformers == 0:
                    row = table.rows[0].cells
                    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                elif four_conformers > 0 and p == 1:
                    row = table.add_row().cells
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("Conformer " + str(conformer_number + four_conformers * 4 + 1))
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(11)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row[p].merge(row[p + 1])
                row[p].merge(row[p + 2])
            # Add conformer energies

            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if settings["Boltz energy type"] == "Gibbs free energy":
                run = paragraph.add_run("Gibbs free energy (hartrees)")
            elif settings["Boltz energy type"] == "Electronic energy":
                run = paragraph.add_run("Electronic energy (hartrees)")
            run.font.size = Pt(8)
            run.font.bold = True
            set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run(str(round(table_energies[four_conformers][conformer_number], 8)))
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row[p].merge(row[p + 1])
                row[p].merge(row[p + 2])
            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if settings["Relative energy unit"] == "kcal/mol" and settings["Boltz energy type"] == "Gibbs free energy":
                run = paragraph.add_run("ΔG (kcal/mol)")
            elif settings["Relative energy unit"] == "kJ/mol" and settings["Boltz energy type"] == "Gibbs free energy":
                run = paragraph.add_run("ΔG (kJ/mol)")
            elif (
                settings["Relative energy unit"] == "kcal/mol" and settings["Boltz energy type"] == "Electronic energy"
            ):
                run = paragraph.add_run("ΔE (kcal/mol)")
            elif settings["Relative energy unit"] == "kJ/mol" and settings["Boltz energy type"] == "Electronic energy":
                run = paragraph.add_run("ΔE (kJ/mol)")
            run.font.size = Pt(8)
            run.font.bold = True
            set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    if conformer_number == 0 and four_conformers == 0:
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(round(table_relative_energies[four_conformers][conformer_number], 3))
                        )
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row[p].merge(row[p + 1])
                row[p].merge(row[p + 2])
            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("Population Proportion")
            run.font.size = Pt(8)
            run.font.bold = True
            set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run(
                        str(round((table_boltz_weights[four_conformers][conformer_number] * 100), 2)) + "%"
                    )
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row[p].merge(row[p + 1])
                row[p].merge(row[p + 2])
            # Add column headings

            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("Element")
            run.font.size = Pt(8)
            run.font.bold = True
            set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("X")
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p += 1
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("Y")
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p += 1
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("Z")
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(8)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Fill table with conformer Cartesian coordinates

            for atom in range(0, len(ordered_element_list[0])):
                row = table.add_row().cells
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = row[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(ordered_element_list[0][atom])
                run.font.size = Pt(8)
                if atom == len(ordered_element_list[0]) - 1:
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                for conformer_number in table_loop:
                    r = 1 + 3 * conformer_number
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border, if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(
                                f"{round(float(table_x_cartesian_coords_list[four_conformers][conformer_number][atom]), 6):6f}"
                            )
                        )
                        if atom == len(ordered_element_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(6.5)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    r += 1
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border, if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(
                                f"{round(float(table_y_cartesian_coords_list[four_conformers][conformer_number][atom]), 6):6f}"
                            )
                        )
                        if atom == len(ordered_element_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(6.5)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    r += 1
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border, if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(
                                f"{round(float(table_z_cartesian_coords_list[four_conformers][conformer_number][atom]), 6):6f}"
                            )
                        )
                        if atom == len(ordered_element_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(6.5)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if settings["NMR/ECD/VCD/OR table"] is False:
        doc_name = results_name_and_directory + " Supplementary Tables.docx"
        # Save the Word document.

        try:
            doc.save(doc_name)
            return "", ""
        except:
            return (
                "Error detected",
                "Unable to save Microsoft Word document.\n Try closing this Word document, then drag and drop your "
                "files again.",
            )
    # Create table for ECD data, if present

    if ordered_wavelength_list:
        if settings["Energies and coordinates table"] is False:
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            # Create style

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1
        # Add ECD table title

        if doc.tables:
            doc.add_page_break()
        if not tddft_solvent:
            tddft_solvation_text = " in the gas phase"
        else:
            tddft_solvation_text = " in " + tddft_solvent
        table_title = doc.add_heading("", level=2)

        if settings["UV in ECD table"] is False:
            table_title.add_run(
                "Calculated excited state transition wavelengths and rotatory strengths for geometry-optimized "
                "conformers of "
                + results_name
                + " at "
                + tddft_functional_and_basis_set
                + tddft_solvation_text
                + "\n"
            )
            # Create ECD table

            number_of_columns = 1 + 2 * len(table_wavelength_list[0])
            number_existing_tables = len(doc.tables)
            for four_conformers in range(0, len(table_wavelength_list)):
                table = doc.add_table(1, number_of_columns)
                set_cell_margins(table)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Remove extra first row if is concatenated table

                if four_conformers > 0:
                    table_number = number_existing_tables + four_conformers
                    Table = doc.tables[table_number]
                    RowA = Table.rows[0]
                    table_element = Table._tbl
                    table_element.remove(RowA._tr)
                # Add conformer numbers

                for conformer_number in range(0, len(table_wavelength_list[four_conformers])):
                    p = 1 + conformer_number * 2
                    if four_conformers == 0:
                        row = table.rows[0].cells
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    elif four_conformers > 0 and p == 1:
                        row = table.add_row().cells
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    p = 1 + conformer_number * 2
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("Conformer " + str(conformer_number + four_conformers * 4 + 1))
                    run.font.size = Pt(11)
                    run.font.bold = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    row[p].merge(row[p + 1])
                # Add column headings

                row = table.add_row().cells
                paragraph = row[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("State")
                run.font.size = Pt(11)
                run.font.bold = True
                set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for conformer_number in range(0, len(table_wavelength_list[four_conformers])):
                    p = 1 + conformer_number * 2
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border.

                    if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run("λ")
                        set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(11)
                    run.font.bold = True
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p += 1
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border.

                    if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run("𝑅")
                        set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(11)
                    run.font.bold = True
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Fill ECD table with wavelengths and rotatory strengths

                if len(table_energies) > 1:
                    table_loop = range(4)
                else:
                    table_loop = range(len(table_wavelength_list[four_conformers]))
                for excited_state_number, excited_state in enumerate(ordered_wavelength_list[0]):
                    row = table.add_row().cells
                    paragraph = row[0].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run(str(excited_state_number + 1))
                    run.font.size = Pt(10)
                    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    for conformer_number in table_loop:
                        r = 1 + 2 * conformer_number
                        paragraph = row[r].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # If we have run out of conformers, fill cell with a single space and no bottom border.

                        if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                            run = paragraph.add_run(" ")
                        else:
                            run = paragraph.add_run(
                                table_wavelength_list[four_conformers][conformer_number][excited_state_number]
                            )
                        run.font.size = Pt(10)
                        row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                        r += 1
                        paragraph = row[r].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # If we have run out of conformers, fill cell with a single space and no bottom border.
                        if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                            run = paragraph.add_run(" ")
                        else:
                            run = paragraph.add_run(
                                str(
                                    round(
                                        float(
                                            table_rotatory_strength_list[four_conformers][conformer_number][
                                                excited_state_number
                                            ]
                                        ),
                                        4,
                                    )
                                )
                            )
                        run.font.size = Pt(10)
                        row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            doc.add_paragraph("λ = wavelength (nm), 𝑅 = rotatory strength (× 10⁻⁴⁰ esu cm erg/G).")
            # Create text for name for Word document.

            doc_text += "ECD"
        else:
            table_title.add_run(
                "Calculated excited state transition wavelengths, oscillator strengths and rotatory strengths for "
                "geometry-optimized conformers of "
                + results_name
                + " at "
                + tddft_functional_and_basis_set
                + tddft_solvation_text
                + "\n"
            )
            # Create ECD/UV table

            number_of_columns = 1 + 3 * len(table_wavelength_list[0])
            number_existing_tables = len(doc.tables)
            for four_conformers in range(0, len(table_wavelength_list)):
                table = doc.add_table(1, number_of_columns)
                set_cell_margins(table)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                # Remove extra first row if is concatenated table

                if four_conformers > 0:
                    table_number = number_existing_tables + four_conformers
                    Table = doc.tables[table_number]
                    RowA = Table.rows[0]
                    table_element = Table._tbl
                    table_element.remove(RowA._tr)
                # Add conformer numbers

                for conformer_number in range(0, len(table_wavelength_list[four_conformers])):
                    p = 1 + conformer_number * 3
                    if four_conformers == 0:
                        row = table.rows[0].cells
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    elif four_conformers > 0 and p == 1:
                        row = table.add_row().cells
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    p = 1 + conformer_number * 3
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("Conformer " + str(conformer_number + four_conformers * 4 + 1))
                    run.font.size = Pt(11)
                    run.font.bold = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    row[p].merge(row[p + 1])
                    row[p].merge(row[p + 2])
                # Add column headings

                row = table.add_row().cells
                paragraph = row[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("State")
                run.font.size = Pt(11)
                run.font.bold = True
                set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for conformer_number in range(0, len(table_wavelength_list[four_conformers])):
                    p = 1 + conformer_number * 3
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("λ")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p += 1
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("𝑓")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    p += 1
                    paragraph = row[p].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run("𝑅")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # Fill ECD/UV table with wavelengths, oscillator strengths and rotatory strengths

                if len(table_energies) > 1:
                    table_loop = range(4)
                else:
                    table_loop = range(len(table_wavelength_list[four_conformers]))
                for excited_state_number, excited_state in enumerate(ordered_wavelength_list[0]):
                    row = table.add_row().cells
                    paragraph = row[0].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run(str(excited_state_number + 1))
                    run.font.size = Pt(9)
                    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                        set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    for conformer_number in table_loop:
                        r = 1 + 3 * conformer_number
                        paragraph = row[r].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # If we have run out of conformers, fill cell with a single space and no bottom border.

                        if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                            run = paragraph.add_run(" ")
                        else:
                            run = paragraph.add_run(
                                table_wavelength_list[four_conformers][conformer_number][excited_state_number]
                            )
                        run.font.size = Pt(8)
                        row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                        r += 1
                        paragraph = row[r].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # If we have run out of conformers, fill cell with a single space and no bottom border.

                        if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                            run = paragraph.add_run(" ")
                        else:
                            run = paragraph.add_run(
                                str(
                                    round(
                                        float(
                                            table_oscillator_strength_list[four_conformers][conformer_number][
                                                excited_state_number
                                            ]
                                        ),
                                        4,
                                    )
                                )
                            )
                        run.font.size = Pt(8)
                        row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                        r += 1
                        paragraph = row[r].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # If we have run out of conformers, fill cell with a single space and no bottom border.

                        if conformer_number > (len(table_wavelength_list[four_conformers]) - 1):
                            run = paragraph.add_run(" ")
                        else:
                            run = paragraph.add_run(
                                str(
                                    round(
                                        float(
                                            table_rotatory_strength_list[four_conformers][conformer_number][
                                                excited_state_number
                                            ]
                                        ),
                                        4,
                                    )
                                )
                            )
                        run.font.size = Pt(8)
                        row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        if excited_state_number == len(ordered_wavelength_list[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            doc.add_paragraph(
                "λ = wavelength (nm), 𝑓 = oscillator strength, 𝑅 = rotatory strength (× 10⁻⁴⁰ esu cm erg/G).")
            # Create text for name for Word document.

            doc_text += "ECD+UV"

    def six_conformer_table(data_type, first_column_data, table_data, functional_and_basis_set, solvent, results_name):
        """Writes a data table containing either NMR shielding tensors or optical rotations, with six conformers
        per row."""

        # Add table title

        if doc.tables:
            doc.add_page_break()
        if not solvent:
            solvation_text = " in the gas phase"
        else:
            solvation_text = " in " + solvent
        table_title = doc.add_heading("", level=2)
        if data_type == "nmr":
            table_title.add_run(
                "Calculated NMR isotropic shielding tensors for geometry-optimized conformers of "
                + results_name
                + " at "
                + functional_and_basis_set
                + solvation_text
                + "\n"
            )
            first_column_heading = "Atom"
            first_column = 1
        elif data_type == "or":
            table_title.add_run(
                "Calculated specific rotations for geometry-optimized conformers of "
                + results_name
                + " at "
                + functional_and_basis_set
                + solvation_text
                + "\n"
            )
            first_column_heading = "λ (nm)"
            first_column = 1
        # Create table

        number_of_columns = first_column + 1 * len(table_data[0])
        number_existing_tables = len(doc.tables)
        for six_conformers in range(0, len(table_data)):
            table = doc.add_table(1, number_of_columns)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Remove extra first row if is concatenated table

            if six_conformers > 0:
                table_number = number_existing_tables + six_conformers
                Table = doc.tables[table_number]
                RowA = Table.rows[0]
                table_element = Table._tbl
                table_element.remove(RowA._tr)
            # Add column headings

            for conformer_number in range(0, len(table_data[six_conformers])):
                p = first_column + conformer_number
                if six_conformers == 0:
                    row = table.rows[0].cells
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                elif six_conformers > 0 and p == 1:
                    row = table.add_row().cells
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                if conformer_number == 0:
                    if first_column == 1:
                        run = row[0].paragraphs[0].add_run(first_column_heading)
                        run.font.size = Pt(11)
                        run.font.bold = True
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = row[0].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p = first_column + conformer_number
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run("Conformer " + str(conformer_number + six_conformers * 6 + 1))
                run.font.size = Pt(11)
                run.font.bold = True
                set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Fill table with shielding tensors

            for data_entry in range(0, len(table_data[0][0])):
                row = table.add_row().cells
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = row[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if first_column == 1:
                    run = paragraph.add_run(first_column_data[0][data_entry])
                    run.font.size = Pt(11)
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if data_entry == len(table_data[0][0]) - 1:
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                for conformer_number in range(0, len(table_data[six_conformers])):
                    r = first_column + conformer_number
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.add_run(str(table_data[six_conformers][conformer_number][data_entry]))
                    run.font.size = Pt(11)
                    row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if data_entry == len(table_data[0][0]) - 1:
                        set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})

    # Create table for NMR data, if present

    if ordered_shielding_tensors:
        if settings["Energies and coordinates table"] is False:
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            # Create style

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1
        six_conformer_table(
            "nmr", ordered_element_list, table_shielding_tensors, nmr_functional_and_basis_set, nmr_solvent,
            results_name
        )
        doc_text += "+NMR"
    # Create table for VCD data, if present

    calc_software = analysed_data[55]
    if ordered_frequency_rotatory_strengths_list:
        if settings["Energies and coordinates table"] is False:
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            # Create style

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1
        # Add table title

        number_existing_tables = len(doc.tables)
        if number_existing_tables > 0:
            doc.add_page_break()
        if not opt_freq_solvent:
            solvation_text = " in the gas phase"
        else:
            solvation_text = " in " + opt_freq_solvent
        table_title = doc.add_heading("", level=2)
        if calc_software == "gaussian":
            table_title.add_run(
                "Calculated vibrational frequencies, dipole and rotational strengths for geometry-optimized "
                "conformers of "
                + results_name
                + " at "
                + opt_freq_functional_and_basis_set
                + solvation_text
                + "\n"
            )
        elif calc_software == "orca":
            table_title.add_run(
                "Calculated vibrational frequencies, molar extinction coefficients and rotational strengths for "
                "geometry-optimized conformers of "
                + results_name
                + " at "
                + opt_freq_functional_and_basis_set
                + solvation_text
                + "\n"
            )
        first_column_heading = "Atom"
        first_column = 1

        # Create style

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1

        number_of_columns = 1 + 3 * len(table_energies[0])
        # Create table

        for four_conformers in range(0, len(table_energies)):
            table = doc.add_table(1, number_of_columns)
            set_cell_margins(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Remove extra first row if is concatenated table

            if four_conformers > 0:
                table_number = number_existing_tables + four_conformers
                Table = doc.tables[table_number]
                RowA = Table.rows[0]
                table_element = Table._tbl
                table_element.remove(RowA._tr)
            # Add conformer numbers

            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                if four_conformers == 0:
                    row = table.rows[0].cells
                    row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                elif four_conformers > 0 and p == 1:
                    row = table.add_row().cells
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("Conformer " + str(conformer_number + four_conformers * 4 + 1))
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(11)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row[p].merge(row[p + 1])
                row[p].merge(row[p + 2])
            # Add column headings

            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("Mode")
            run.font.size = Pt(11)
            run.font.bold = True
            set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
            row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for conformer_number in table_loop:
                p = 1 + conformer_number * 3
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("ν")
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(11)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p += 1
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    if calc_software == "gaussian":
                        run = paragraph.add_run("D")
                    elif calc_software == "orca":
                        run = paragraph.add_run("ε")
                    run.italic = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(11)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p += 1
                paragraph = row[p].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # If we have run out of conformers, fill cell with a single space and no bottom border.

                if conformer_number > (len(table_energies[four_conformers]) - 1):
                    run = paragraph.add_run(" ")
                else:
                    run = paragraph.add_run("R")
                    run.italic = True
                    set_cell_border(row[p], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                run.font.size = Pt(11)
                run.font.bold = True
                row[p].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # Fill table with IR/VCD data

            for mode in range(0, len(ordered_frequencies[0])):
                row = table.add_row().cells
                row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = row[0].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run(str(range(0, len(ordered_frequencies[0]))[mode] + 1))
                run.font.size = Pt(9)
                if mode == len(ordered_frequencies[0]) - 1:
                    set_cell_border(row[0], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                for conformer_number in table_loop:
                    r = 1 + 3 * conformer_number
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border,
                    # if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(round(float(table_frequencies[four_conformers][conformer_number][mode])))
                        )
                        if mode == len(ordered_frequencies[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(9)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    r += 1
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border,
                    # if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        if calc_software == "gaussian":
                            run = paragraph.add_run(
                                str(f"{round(float(table_frequency_dipole_strengths[four_conformers][conformer_number][mode]), 1):.1f}")
                            )
                        elif calc_software == "orca":
                            run = paragraph.add_run(
                                str(f"{round(float(table_frequency_dipole_strengths[four_conformers][conformer_number][mode]), 4):.4f}")
                            )

                        if mode == len(ordered_frequencies[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(9)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    r += 1
                    paragraph = row[r].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    # If we have run out of conformers, fill cell with a single space and no bottom border,
                    # if applicable.

                    if conformer_number > (len(table_energies[four_conformers]) - 1):
                        run = paragraph.add_run(" ")
                    else:
                        run = paragraph.add_run(
                            str(
                                f"{round(float(table_frequency_rotatory_strengths[four_conformers][conformer_number][mode]), 1):.1f}"
                            )
                        )
                        if mode == len(ordered_frequencies[0]) - 1:
                            set_cell_border(row[r], bottom={"sz": 4, "val": "single", "color": "auto", "space": "0"})
                    run.font.size = Pt(9)
                    row[r].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if calc_software == "gaussian":
            doc.add_paragraph(
                "ν = frequency (cm⁻¹), 𝐷 = dipole strength (× 10⁻⁴⁰ esu² cm²), 𝑅 = rotational strength (× 10⁻⁴⁴ esu² cm²)"
            )
        elif calc_software == "orca":
            doc.add_paragraph(
                "ν = frequency (cm⁻¹), 𝜀 = molar extinction coefficient (L mol⁻¹ cm⁻¹), 𝑅 = rotational strength (× 10⁻⁴⁴ esu² cm²)"
            )
        doc_text += "+VCD"
    # Create table for optrot data, if present

    if ordered_optrot_wavelengths_list:
        if settings["Energies and coordinates table"] is False:
            doc = Document()
            doc.core_properties.author = 'SpectroIBIS'
            # Create style

            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1
        six_conformer_table(
            "or", ordered_optrot_wavelengths_list, table_optrots, or_functional_and_basis_set, or_solvent, results_name
        )
        doc_text += "+OR"
    if (
            0
            == len(table_wavelength_list)
            == len(ordered_shielding_tensors)
            == len(ordered_frequency_rotatory_strengths_list)
            == len(ordered_optrot_wavelengths_list)
    ):
        doc_text += "+Energies + Coordinates Only"
    # Make filename

    doc_text = doc_text.lstrip("+")
    doc_name = results_name_and_directory + " Supplementary Tables (" + doc_text + ").docx"
    # Save the Word document.

    try:
        doc.save(doc_name)
        return "", ""
    except:
        return (
            "Error detected",
            "Unable to save Microsoft Word document.\n Try closing this Word document, then drag and drop your files "
            "again.",
        )
